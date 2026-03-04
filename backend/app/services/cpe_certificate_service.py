"""CPE Certificate service for eligibility checking, issuance, and PDF generation."""
import asyncio
import io
import logging
from datetime import datetime, timedelta, timezone
from typing import Optional

from sqlalchemy import select, func, and_
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.models.audit_log import AuditLog
from app.models.cpe_certificate import CPECertificate, CertificateStatus
from app.models.event import Event, EventParticipation, ParticipationStatus
from app.models.user import User
from app.models.vpn import VPNCredential

logger = logging.getLogger(__name__)

# Module-level caches (downloaded from R2 once per process)
_template_cache: Optional[bytes] = None
_signature_1_cache: Optional[bytes] = None
_signature_2_cache: Optional[bytes] = None


class CPECertificateService:
    """Service for CPE certificate eligibility, issuance, and PDF generation."""

    def __init__(self, session: AsyncSession):
        self.session = session
        self.settings = get_settings()

    async def check_eligibility(self, user_id: int, event: Event) -> dict:
        """
        Check a single user's CPE eligibility for an event.

        Criteria (all required for 32 CPE hours):
        1. At least one Nextcloud login during event dates
        2. At least one PowerDNS-Admin login during event dates
        3. At least one VPN credential assigned

        Returns dict with eligible flag, criteria details, and CPE hours.
        """
        event_start = datetime.combine(event.start_date, datetime.min.time()).replace(tzinfo=timezone.utc)
        event_end = datetime.combine(event.end_date, datetime.min.time()).replace(tzinfo=timezone.utc) + timedelta(days=1)

        has_nextcloud = await self._check_login(user_id, "nextcloud", event_start, event_end)
        has_powerdns = await self._check_login(user_id, "powerdns-admin", event_start, event_end)
        has_vpn = await self._check_vpn_assigned(user_id)

        eligible = has_nextcloud and has_powerdns and has_vpn

        return {
            "user_id": user_id,
            "eligible": eligible,
            "cpe_hours": self.settings.CPE_HOURS_DEFAULT if eligible else 0,
            "criteria": {
                "has_nextcloud_login": has_nextcloud,
                "has_powerdns_login": has_powerdns,
                "has_vpn_assigned": has_vpn,
            }
        }

    async def bulk_check_eligibility(self, event: Event) -> list[dict]:
        """
        Check CPE eligibility for all confirmed participants of an event.

        Returns list of eligibility results, one per confirmed participant.
        """
        # Get all confirmed participants
        result = await self.session.execute(
            select(User)
            .join(EventParticipation, and_(
                EventParticipation.user_id == User.id,
                EventParticipation.event_id == event.id
            ))
            .where(
                EventParticipation.status == ParticipationStatus.CONFIRMED.value,
                User.is_active == True,
            )
            .order_by(User.last_name, User.first_name)
        )
        users = result.scalars().all()

        results = []
        for user in users:
            eligibility = await self.check_eligibility(user.id, event)
            eligibility["email"] = user.email
            eligibility["first_name"] = user.first_name
            eligibility["last_name"] = user.last_name

            # Check if certificate already exists
            existing = await self._get_existing_certificate(user.id, event.id)
            eligibility["certificate_exists"] = existing is not None
            if existing:
                eligibility["certificate_number"] = existing.certificate_number
                eligibility["certificate_status"] = existing.status

            results.append(eligibility)

        return results

    async def issue_certificate(
        self, user_id: int, event_id: int, issued_by_user_id: int,
        skip_eligibility: bool = False,
    ) -> CPECertificate:
        """
        Issue a CPE certificate to a user for an event.

        Verifies eligibility (unless skip_eligibility=True), generates certificate
        number, creates DB record, generates DOCX, converts to PDF, and uploads to R2.

        Raises ValueError if user is ineligible or certificate already exists.
        """
        # Load event
        event = await self.session.get(Event, event_id)
        if not event:
            raise ValueError(f"Event {event_id} not found")

        # Check for existing certificate
        existing = await self._get_existing_certificate(user_id, event_id)
        if existing:
            if existing.status == CertificateStatus.REVOKED.value:
                # Delete revoked cert so a fresh one can be issued
                logger.info(
                    f"Removing revoked certificate {existing.certificate_number} "
                    f"for user {user_id} to allow re-issuance"
                )
                await self.session.delete(existing)
                await self.session.flush()
            else:
                raise ValueError(
                    f"Certificate already exists: {existing.certificate_number} "
                    f"(status: {existing.status})"
                )

        # Check eligibility (still record snapshot even when skipping)
        eligibility = await self.check_eligibility(user_id, event)
        if not eligibility["eligible"] and not skip_eligibility:
            missing = [k for k, v in eligibility["criteria"].items() if not v]
            raise ValueError(f"User {user_id} is ineligible. Missing: {', '.join(missing)}")

        # Load user
        user = await self.session.get(User, user_id)
        if not user:
            raise ValueError(f"User {user_id} not found")

        # Generate certificate number
        cert_number = await self._generate_certificate_number(event.year)

        # Create certificate record
        certificate = CPECertificate(
            user_id=user_id,
            event_id=event_id,
            issued_by_user_id=issued_by_user_id,
            certificate_number=cert_number,
            cpe_hours=self.settings.CPE_HOURS_DEFAULT,
            status=CertificateStatus.ISSUED.value,
            has_nextcloud_login=eligibility["criteria"]["has_nextcloud_login"],
            has_powerdns_login=eligibility["criteria"]["has_powerdns_login"],
            has_vpn_assigned=eligibility["criteria"]["has_vpn_assigned"],
        )
        self.session.add(certificate)
        await self.session.flush()

        # Generate and upload PDF
        try:
            pdf_storage_key = await self._generate_and_upload_pdf(certificate, user, event)
            certificate.pdf_storage_key = pdf_storage_key
            certificate.pdf_generated_at = datetime.now(timezone.utc)
        except Exception as e:
            logger.error(f"Failed to generate PDF for {cert_number}: {e}")
            # Certificate is still created but without PDF - can be regenerated later

        await self.session.flush()
        return certificate

    async def bulk_issue(
        self, event_id: int, user_ids: list[int], issued_by_user_id: int,
        skip_eligibility: bool = False,
    ) -> dict:
        """
        Issue certificates in bulk for multiple users.

        Returns summary with lists of issued, skipped_ineligible,
        skipped_existing, and failed user IDs.
        """
        event = await self.session.get(Event, event_id)
        if not event:
            raise ValueError(f"Event {event_id} not found")

        issued = []
        skipped_ineligible = []
        skipped_existing = []
        failed = []

        for user_id in user_ids:
            try:
                existing = await self._get_existing_certificate(user_id, event_id)
                if existing and existing.status != CertificateStatus.REVOKED.value:
                    skipped_existing.append({
                        "user_id": user_id,
                        "certificate_number": existing.certificate_number,
                    })
                    continue

                if not skip_eligibility:
                    eligibility = await self.check_eligibility(user_id, event)
                    if not eligibility["eligible"]:
                        skipped_ineligible.append({
                            "user_id": user_id,
                            "criteria": eligibility["criteria"],
                        })
                        continue

                cert = await self.issue_certificate(
                    user_id, event_id, issued_by_user_id,
                    skip_eligibility=skip_eligibility,
                )
                issued.append({
                    "user_id": user_id,
                    "certificate_number": cert.certificate_number,
                })

            except Exception as e:
                logger.error(f"Failed to issue certificate for user {user_id}: {e}")
                failed.append({"user_id": user_id, "error": str(e)})

        await self.session.commit()

        return {
            "issued": issued,
            "skipped_ineligible": skipped_ineligible,
            "skipped_existing": skipped_existing,
            "failed": failed,
        }

    async def revoke_certificate(
        self, certificate_id: int, revoked_by_user_id: int, reason: str
    ) -> CPECertificate:
        """Revoke an issued certificate."""
        cert = await self.session.get(CPECertificate, certificate_id)
        if not cert:
            raise ValueError(f"Certificate {certificate_id} not found")
        if cert.status == CertificateStatus.REVOKED.value:
            raise ValueError(f"Certificate {cert.certificate_number} is already revoked")

        cert.status = CertificateStatus.REVOKED.value
        cert.revoked_at = datetime.now(timezone.utc)
        cert.revoked_by_user_id = revoked_by_user_id
        cert.revocation_reason = reason

        await self.session.flush()
        return cert

    async def regenerate_pdf(self, certificate_id: int) -> CPECertificate:
        """Regenerate the PDF for an existing certificate."""
        cert = await self.session.get(CPECertificate, certificate_id)
        if not cert:
            raise ValueError(f"Certificate {certificate_id} not found")

        user = await self.session.get(User, cert.user_id)
        event = await self.session.get(Event, cert.event_id)

        pdf_storage_key = await self._generate_and_upload_pdf(cert, user, event)
        cert.pdf_storage_key = pdf_storage_key
        cert.pdf_generated_at = datetime.now(timezone.utc)

        await self.session.flush()
        return cert

    async def bulk_regenerate_pdfs(
        self, event_id: int, certificate_ids: Optional[list[int]] = None
    ) -> dict:
        """
        Regenerate PDFs for certificates missing them.

        If certificate_ids provided, regenerate those specific certs.
        If only event_id, regenerate all ISSUED certs missing a PDF.
        """
        if certificate_ids:
            query = (
                select(CPECertificate)
                .where(
                    CPECertificate.id.in_(certificate_ids),
                    CPECertificate.event_id == event_id,
                )
            )
        else:
            query = (
                select(CPECertificate)
                .where(
                    CPECertificate.event_id == event_id,
                    CPECertificate.status == CertificateStatus.ISSUED.value,
                    CPECertificate.pdf_storage_key.is_(None),
                )
            )

        result = await self.session.execute(query)
        certs = result.scalars().all()

        regenerated = []
        failed = []
        skipped_revoked = []

        for cert in certs:
            if cert.status == CertificateStatus.REVOKED.value:
                skipped_revoked.append({
                    "certificate_id": cert.id,
                    "certificate_number": cert.certificate_number,
                })
                continue

            try:
                user = await self.session.get(User, cert.user_id)
                event = await self.session.get(Event, cert.event_id)

                pdf_storage_key = await self._generate_and_upload_pdf(cert, user, event)
                cert.pdf_storage_key = pdf_storage_key
                cert.pdf_generated_at = datetime.now(timezone.utc)
                await self.session.flush()

                regenerated.append({
                    "certificate_id": cert.id,
                    "certificate_number": cert.certificate_number,
                })
                logger.info(f"Regenerated PDF for {cert.certificate_number}")
            except Exception as e:
                logger.error(f"Failed to regenerate PDF for {cert.certificate_number}: {e}")
                failed.append({
                    "certificate_id": cert.id,
                    "certificate_number": cert.certificate_number,
                    "error": str(e),
                })

        return {
            "regenerated": regenerated,
            "failed": failed,
            "skipped_revoked": skipped_revoked,
        }

    async def get_certificates_for_event(
        self, event_id: int, status: Optional[str] = None
    ) -> list[CPECertificate]:
        """Get all certificates for an event, optionally filtered by status."""
        query = (
            select(CPECertificate)
            .where(CPECertificate.event_id == event_id)
            .order_by(CPECertificate.certificate_number)
        )
        if status:
            query = query.where(CPECertificate.status == status)

        result = await self.session.execute(query)
        return result.scalars().all()

    async def get_user_certificates(self, user_id: int) -> list[CPECertificate]:
        """Get all certificates for a user."""
        result = await self.session.execute(
            select(CPECertificate)
            .where(
                CPECertificate.user_id == user_id,
                CPECertificate.status == CertificateStatus.ISSUED.value,
            )
            .order_by(CPECertificate.created_at.desc())
        )
        return result.scalars().all()

    async def get_certificate_by_id(self, certificate_id: int) -> Optional[CPECertificate]:
        """Get a certificate by ID."""
        return await self.session.get(CPECertificate, certificate_id)

    # -------------------------------------------------------------------------
    # Private helpers
    # -------------------------------------------------------------------------

    async def _check_login(
        self, user_id: int, client_id: str,
        event_start: datetime, event_end: datetime
    ) -> bool:
        """Check if user has at least one Keycloak login for a specific client during event dates."""
        result = await self.session.execute(
            select(func.count(AuditLog.id)).where(
                AuditLog.user_id == user_id,
                AuditLog.action == "KEYCLOAK_LOGIN",
                AuditLog.details["client_id"].as_string() == client_id,
                AuditLog.created_at >= event_start,
                AuditLog.created_at <= event_end,
            )
        )
        count = result.scalar()
        return count > 0

    async def _check_vpn_assigned(self, user_id: int) -> bool:
        """Check if user has at least one VPN credential assigned."""
        result = await self.session.execute(
            select(func.count(VPNCredential.id)).where(
                VPNCredential.assigned_to_user_id == user_id
            )
        )
        count = result.scalar()
        return count > 0

    async def _get_existing_certificate(
        self, user_id: int, event_id: int
    ) -> Optional[CPECertificate]:
        """Check if a certificate already exists for this user+event."""
        result = await self.session.execute(
            select(CPECertificate).where(
                CPECertificate.user_id == user_id,
                CPECertificate.event_id == event_id,
            )
        )
        return result.scalar_one_or_none()

    async def _generate_certificate_number(self, year: int) -> str:
        """Generate a unique certificate number with random hex serial (CX-YYYY-XXXX)."""
        import secrets

        for _ in range(10):
            serial = secrets.token_hex(2).upper()  # 4 hex chars, 65536 possibilities
            cert_number = f"CX-{year}-{serial}"

            # Check uniqueness
            result = await self.session.execute(
                select(func.count(CPECertificate.id)).where(
                    CPECertificate.certificate_number == cert_number
                )
            )
            if result.scalar() == 0:
                return cert_number

        raise RuntimeError(f"Failed to generate unique certificate number after 10 attempts")

    async def _generate_and_upload_pdf(
        self, certificate: CPECertificate, user: User, event: Event
    ) -> str:
        """
        Generate a filled DOCX from the template, convert to PDF,
        overlay signature image, and upload to R2. Returns the R2 storage key.
        """
        # Fetch template from R2 (cached after first download)
        template_bytes = self._get_template_bytes()

        # Fill the DOCX template (text only — signatures handled as PDF overlay)
        docx_bytes = self._fill_template(certificate, user, event, template_bytes)

        # Convert DOCX to PDF via Gotenberg/LibreOffice
        pdf_bytes = await self._convert_to_pdf(docx_bytes)

        # Overlay signature images onto the PDF
        sig1_bytes = self._get_signature_image(
            self.settings.CPE_SIGNATURE_1_R2_KEY, "_signature_1_cache"
        )
        sig2_bytes = self._get_signature_image(
            self.settings.CPE_SIGNATURE_2_R2_KEY, "_signature_2_cache"
        )
        if sig1_bytes or sig2_bytes:
            pdf_bytes = self._overlay_signatures(pdf_bytes, sig1_bytes, sig2_bytes)

        # Upload to R2
        storage_key = f"certificates/{event.slug}/{certificate.certificate_number}.pdf"
        self._upload_to_r2(storage_key, pdf_bytes)

        logger.info(
            f"Generated and uploaded PDF for {certificate.certificate_number} "
            f"({len(pdf_bytes)} bytes)"
        )
        return storage_key

    def _get_template_bytes(self) -> bytes:
        """Download the DOCX template from R2, caching in memory after first fetch."""
        global _template_cache

        if _template_cache is not None:
            return _template_cache

        r2_key = self.settings.CPE_TEMPLATE_R2_KEY
        if not r2_key:
            raise ValueError("CPE_TEMPLATE_R2_KEY is not configured")

        import boto3
        from botocore.config import Config

        account_id = self.settings.R2_ACCOUNT_ID
        access_key_id = self.settings.R2_ACCESS_KEY_ID
        secret_access_key = self.settings.R2_SECRET_ACCESS_KEY
        bucket = self.settings.R2_BUCKET

        if not all([account_id, access_key_id, secret_access_key, bucket]):
            raise ValueError(
                "R2 must be configured to fetch the CPE template "
                "(R2_ACCOUNT_ID, R2_ACCESS_KEY_ID, R2_SECRET_ACCESS_KEY, R2_BUCKET)"
            )

        endpoint_url = f"https://{account_id}.r2.cloudflarestorage.com"

        s3 = boto3.client(
            "s3",
            endpoint_url=endpoint_url,
            aws_access_key_id=access_key_id,
            aws_secret_access_key=secret_access_key,
            region_name="auto",
            config=Config(signature_version="s3v4"),
        )

        response = s3.get_object(Bucket=bucket, Key=r2_key)
        _template_cache = response["Body"].read()

        logger.info(f"Downloaded CPE template from R2: {r2_key} ({len(_template_cache)} bytes)")
        return _template_cache

    def _get_signature_image(self, r2_key: str, cache_name: str) -> Optional[bytes]:
        """Download a signature image from R2, caching in memory after first fetch."""
        global _signature_1_cache, _signature_2_cache

        if not r2_key:
            return None

        # Check module-level cache
        cached = globals().get(cache_name)
        if cached is not None:
            return cached

        import boto3
        from botocore.config import Config

        account_id = self.settings.R2_ACCOUNT_ID
        access_key_id = self.settings.R2_ACCESS_KEY_ID
        secret_access_key = self.settings.R2_SECRET_ACCESS_KEY
        bucket = self.settings.R2_BUCKET

        if not all([account_id, access_key_id, secret_access_key, bucket]):
            logger.warning("R2 not configured — cannot fetch signature image")
            return None

        endpoint_url = f"https://{account_id}.r2.cloudflarestorage.com"

        s3 = boto3.client(
            "s3",
            endpoint_url=endpoint_url,
            aws_access_key_id=access_key_id,
            aws_secret_access_key=secret_access_key,
            region_name="auto",
            config=Config(signature_version="s3v4"),
        )

        try:
            response = s3.get_object(Bucket=bucket, Key=r2_key)
            data = response["Body"].read()
            globals()[cache_name] = data
            logger.info(f"Downloaded signature image from R2: {r2_key} ({len(data)} bytes)")
            return data
        except Exception as e:
            logger.error(f"Failed to download signature image from R2: {e}")
            return None

    # Official Canadian flag SVG (Pantone colors) from Wikimedia Commons.
    # Embedded to avoid external downloads at runtime.
    _CANADA_FLAG_SVG = (
        '<svg xmlns="http://www.w3.org/2000/svg" width="1200" height="600" viewBox="0 0 9600 4800">'
        '<path fill="#d52b1e" d="m0 0h2400l99 99h4602l99-99h2400v4800h-2400l-99-99h-4602l-99 99H0z"/>'
        '<path fill="#fff" d="m2400 0h4800v4800h-4800zm2490 4430-45-863a95 95 0 0 1 111-98l859 151-116-320'
        'a65 65 0 0 1 20-73l941-762-212-99a65 65 0 0 1-34-79l186-572-542 115a65 65 0 0 1-73-38l-105-247-423'
        ' 454a65 65 0 0 1-111-57l204-1052-327 189a65 65 0 0 1-91-27l-332-652-332 652a65 65 0 0 1-91 27l-327'
        '-189 204 1052a65 65 0 0 1-111 57l-423-454-105 247a65 65 0 0 1-73 38l-542-115 186 572a65 65 0 0 1-34'
        ' 79l-212 99 941 762a65 65 0 0 1 20 73l-116 320 859-151a95 95 0 0 1 111 98l-45 863z"/>'
        '</svg>'
    )

    @staticmethod
    def _make_canadian_flag(width: int = 300, height: int = 150, alpha: int = 40) -> io.BytesIO:
        """Render the official Canadian flag SVG to a translucent PNG."""
        import cairosvg
        from PIL import Image

        png_bytes = cairosvg.svg2png(
            bytestring=CPECertificateService._CANADA_FLAG_SVG.encode(),
            output_width=width, output_height=height,
        )
        img = Image.open(io.BytesIO(png_bytes)).convert("RGBA")

        # Scale alpha channel to desired opacity
        r, g, b, a = img.split()
        a = a.point(lambda x: int(x * (alpha / 255)))
        img = Image.merge("RGBA", (r, g, b, a))

        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf

    @staticmethod
    def _make_american_flag(width: int = 304, height: int = 160, alpha: int = 40) -> io.BytesIO:
        """Generate a translucent American flag PNG in memory."""
        from PIL import Image, ImageDraw

        img = Image.new("RGBA", (width, height), (0, 0, 0, 0))
        d = ImageDraw.Draw(img)

        stripe_h = height / 13
        for i in range(13):
            color = (191, 10, 48, alpha) if i % 2 == 0 else (255, 255, 255, alpha)
            d.rectangle([0, int(i * stripe_h), width, int((i + 1) * stripe_h)], fill=color)

        canton_w = int(width * 0.4)
        canton_h = int(7 * stripe_h)
        d.rectangle([0, 0, canton_w, canton_h], fill=(0, 40, 104, alpha))

        for row in range(5):
            for col in range(6):
                sx = int(canton_w * (col + 0.5) / 6)
                sy = int(canton_h * (row + 0.5) / 5)
                d.ellipse([sx - 3, sy - 3, sx + 3, sy + 3], fill=(255, 255, 255, alpha))
        for row in range(4):
            for col in range(5):
                sx = int(canton_w * (col + 1) / 6)
                sy = int(canton_h * (row + 1) / 5)
                d.ellipse([sx - 3, sy - 3, sx + 3, sy + 3], fill=(255, 255, 255, alpha))

        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf

    @staticmethod
    def _find_signature_line_y(pdf_bytes: bytes) -> float:
        """Find the y-position (reportlab coords, 0=bottom) of the "Signature" text.

        Uses pdfplumber to extract word positions from the generated PDF,
        so overlays are anchored to the actual rendered layout rather than
        hardcoded coordinates that drift across LibreOffice versions.
        """
        import pdfplumber

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            page = pdf.pages[0]
            page_h = page.height
            for word in page.extract_words():
                if word["text"] == "Signature":
                    # pdfplumber y is top-down; convert to reportlab bottom-up
                    # word["top"] is distance from top of page to top of text
                    sig_y_reportlab = page_h - word["top"]
                    logger.info(
                        f"Found 'Signature' text at pdfplumber y={word['top']:.1f}, "
                        f"reportlab y={sig_y_reportlab:.1f}"
                    )
                    return sig_y_reportlab

        logger.warning("Could not find 'Signature' text in PDF, using fallback y=148")
        return 148.0

    @staticmethod
    def _overlay_signatures(
        pdf_bytes: bytes,
        sig1_bytes: Optional[bytes],
        sig2_bytes: Optional[bytes],
    ) -> bytes:
        """Overlay two signature images with translucent country flags onto the PDF.

        This is done AFTER Gotenberg converts the DOCX to PDF, avoiding all
        LibreOffice rendering issues with inline images in table cells.

        The template is landscape (792x612pt) with a 6-column signature table.
        Signature 1 (left) has a Canadian flag watermark behind it.
        Signature 2 (right) has an American flag watermark behind it.

        Positions are anchored to the "Signature" text detected in the PDF
        via pdfplumber, so they're correct regardless of LibreOffice version.
        """
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.utils import ImageReader
        from pypdf import PdfReader, PdfWriter

        base_reader = PdfReader(io.BytesIO(pdf_bytes))
        page = base_reader.pages[0]
        page_w = float(page.mediabox.width)
        page_h = float(page.mediabox.height)

        # Detect where "Signature" text is in the actual rendered PDF
        sig_text_y = CPECertificateService._find_signature_line_y(pdf_bytes)

        overlay_buf = io.BytesIO()
        c = rl_canvas.Canvas(overlay_buf, pagesize=(page_w, page_h))

        # Signature positions — x and width are fixed (column layout doesn't shift),
        # y is relative to the detected "Signature" text position.
        # "Signature" label is at sig_text_y; the horizontal signature line
        # sits ~8-10pt above that.  Offsets are kept small so signatures
        # don't extend up past the "Program Hours" line.
        sig1_x, sig1_w = 125, 140    # left (Yves)
        sig2_x, sig2_w = 545, 140    # right (Wes)
        sig1_y = sig_text_y - 8      # bottom edge on the signature line
        sig2_y = sig_text_y + 0      # Wes sig is flatter, sits slightly higher

        # Translucent country flags behind signatures
        flag_w = 80
        flag_cy = sig_text_y + 12    # centered behind signature area

        canada_flag = CPECertificateService._make_canadian_flag(alpha=40)
        flag1_img = ImageReader(canada_flag)
        flag1_h = flag_w * 0.5
        flag1_x = sig1_x + (sig1_w - flag_w) / 2 - 10
        c.drawImage(flag1_img, flag1_x, flag_cy - flag1_h / 2,
                     width=flag_w, height=flag1_h, mask="auto")

        usa_flag = CPECertificateService._make_american_flag(alpha=40)
        flag2_img = ImageReader(usa_flag)
        flag2_h = flag_w * 0.527
        flag2_x = sig2_x + (sig2_w - flag_w) / 2 - 10
        c.drawImage(flag2_img, flag2_x, flag_cy - flag2_h / 2,
                     width=flag_w, height=flag2_h, mask="auto")

        # Signature 1 (left — Yves, Canadian)
        if sig1_bytes:
            img1 = ImageReader(io.BytesIO(sig1_bytes))
            img1_w, img1_h = img1.getSize()
            target_h1 = sig1_w * (img1_h / img1_w)
            c.drawImage(img1, sig1_x, sig1_y, width=sig1_w,
                         height=target_h1, mask="auto")

        # Signature 2 (right — Wes, American)
        if sig2_bytes:
            img2 = ImageReader(io.BytesIO(sig2_bytes))
            img2_w, img2_h = img2.getSize()
            target_h2 = sig2_w * (img2_h / img2_w)
            c.drawImage(img2, sig2_x, sig2_y, width=sig2_w,
                         height=target_h2, mask="auto")

        c.save()

        overlay_reader = PdfReader(io.BytesIO(overlay_buf.getvalue()))
        page.merge_page(overlay_reader.pages[0])

        writer = PdfWriter()
        writer.add_page(page)

        output = io.BytesIO()
        writer.write(output)
        result = output.getvalue()

        logger.info(f"Overlaid signature images onto PDF ({len(result)} bytes)")
        return result

    def _fill_template(
        self, certificate: CPECertificate, user: User,
        event: Event, template_bytes: bytes
    ) -> bytes:
        """Fill the DOCX template with certificate data using python-docx.

        Only does text replacement — no image insertion or layout adjustments.
        The signature image is overlaid onto the PDF after conversion.
        """
        from docx import Document

        doc = Document(io.BytesIO(template_bytes))

        # Build replacement map from handlebar template placeholders
        start_str = event.start_date.strftime("%m/%d/%Y") if event.start_date else "TBA"
        end_str = event.end_date.strftime("%m/%d/%Y") if event.end_date else "TBA"
        issue_date_str = certificate.created_at.strftime("%m/%d/%Y") if certificate.created_at else datetime.now(timezone.utc).strftime("%m/%d/%Y")

        replacements = {
            "{{PARTICIPANT_NAME}}": f"{user.first_name} {user.last_name}",
            "{{EVENT_START_DATE}}": start_str,
            "{{EVENT_END_DATE}}": end_str,
            "{{CERT_ID}}": certificate.certificate_number,
            "{{DATE_ISSUED}}": issue_date_str,
            "{{SIGNER_1}}": self.settings.CPE_SIGNER_1_NAME or "",
            "{{SIGNER_2}}": self.settings.CPE_SIGNER_2_NAME or "",
        }

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)

        # Replace in headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header:
                    for paragraph in header.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
            for footer in [section.footer, section.first_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)

        # NOTE: spacing reduction removed — the landscape two-signature template
        # fits on one page without it, and the 30% reduction caused overlay position
        # drift across different LibreOffice versions (local vs Render Gotenberg).

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _reduce_spacing_for_libreoffice(doc):
        """Trim before/after spacing on body paragraphs by 30% for LibreOffice."""
        from docx.oxml.ns import qn

        for p in doc.paragraphs:
            pPr = p._element.find(qn('w:pPr'))
            if pPr is None:
                continue
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                continue
            for attr in (qn('w:after'), qn('w:before')):
                val = spacing.get(attr)
                if val and int(val) > 0:
                    spacing.set(attr, str(int(int(val) * 0.70)))

    @staticmethod
    def _replace_in_paragraph(paragraph, replacements: dict):
        """Replace placeholder text in a paragraph while preserving formatting."""
        full_text = paragraph.text
        if not any(placeholder in full_text for placeholder in replacements):
            return

        for placeholder, replacement in replacements.items():
            if placeholder in full_text:
                # For simple cases where the placeholder is within a single run
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)
                # Re-check full text: if placeholder spans multiple runs,
                # rebuild the paragraph text
                if placeholder in paragraph.text:
                    # Concatenate all runs, replace, and reset
                    combined = "".join(run.text for run in paragraph.runs)
                    new_text = combined.replace(placeholder, replacement)
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                        for run in paragraph.runs[1:]:
                            run.text = ""

    async def _convert_to_pdf(self, docx_bytes: bytes) -> bytes:
        """Convert DOCX bytes to PDF using configured conversion mode."""
        mode = self.settings.CPE_CONVERSION_MODE

        if mode == "gotenberg":
            return await self._convert_via_gotenberg(docx_bytes)
        elif mode == "libreoffice":
            return await self._convert_via_libreoffice(docx_bytes)
        else:
            raise ValueError(f"Unknown CPE_CONVERSION_MODE: {mode}")

    async def _convert_via_gotenberg(self, docx_bytes: bytes) -> bytes:
        """Convert DOCX to PDF via Gotenberg HTTP API.

        Retries on ConnectError to handle the brief window after Render marks
        a deploy as 'live' but the private service isn't fully routable yet.
        """
        import httpx

        gotenberg_url = self.settings.GOTENBERG_URL
        if not gotenberg_url:
            raise ValueError("GOTENBERG_URL is not configured")

        max_retries = 3
        retry_delay = 3  # seconds

        async with httpx.AsyncClient() as client:
            for attempt in range(max_retries + 1):
                try:
                    response = await client.post(
                        f"{gotenberg_url}/forms/libreoffice/convert",
                        files={"files": ("certificate.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                        timeout=30.0,
                    )
                    response.raise_for_status()
                    return response.content
                except httpx.ConnectError:
                    if attempt < max_retries:
                        logger.warning(
                            f"Gotenberg not reachable (attempt {attempt + 1}/{max_retries + 1}), "
                            f"retrying in {retry_delay}s"
                        )
                        await asyncio.sleep(retry_delay)
                    else:
                        logger.error("Gotenberg not reachable after all retries")
                        raise

    async def _convert_via_libreoffice(self, docx_bytes: bytes) -> bytes:
        """Convert DOCX to PDF via local LibreOffice (for self-hosted Docker)."""
        import asyncio
        import tempfile
        import os

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, "certificate.docx")
            with open(docx_path, "wb") as f:
                f.write(docx_bytes)

            process = await asyncio.create_subprocess_exec(
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", tmpdir, docx_path,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            stdout, stderr = await process.communicate()

            if process.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {stderr.decode()}")

            pdf_path = os.path.join(tmpdir, "certificate.pdf")
            with open(pdf_path, "rb") as f:
                return f.read()

    def _upload_to_r2(self, storage_key: str, pdf_bytes: bytes):
        """Upload PDF bytes to Cloudflare R2."""
        import boto3
        from botocore.config import Config

        account_id = self.settings.R2_ACCOUNT_ID
        access_key_id = self.settings.R2_ACCESS_KEY_ID
        secret_access_key = self.settings.R2_SECRET_ACCESS_KEY
        bucket = self.settings.R2_BUCKET

        if not all([account_id, access_key_id, secret_access_key, bucket]):
            raise ValueError(
                "R2 upload requires R2_ACCOUNT_ID, R2_ACCESS_KEY_ID, "
                "R2_SECRET_ACCESS_KEY, and R2_BUCKET to be configured"
            )

        endpoint_url = f"https://{account_id}.r2.cloudflarestorage.com"

        s3 = boto3.client(
            "s3",
            endpoint_url=endpoint_url,
            aws_access_key_id=access_key_id,
            aws_secret_access_key=secret_access_key,
            region_name="auto",
            config=Config(signature_version="s3v4"),
        )

        s3.put_object(
            Bucket=bucket,
            Key=storage_key,
            Body=pdf_bytes,
            ContentType="application/pdf",
        )

        logger.info(f"Uploaded certificate PDF to R2: {storage_key}")
